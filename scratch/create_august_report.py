import docx
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

def create_august_report():
    doc = docx.Document()
    
    # Add title
    p0 = doc.add_paragraph('Monthly Report – August 2026')
    p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.runs[0].bold = True
    p0.runs[0].font.size = Pt(14)
    
    # Add Name
    p1 = doc.add_paragraph('Name: Nurudeen Akindele')
    p1.runs[0].bold = True
    
    # Add Section Title
    p2 = doc.add_paragraph('Daily Task Activities')
    p2.runs[0].bold = True
    p2.runs[0].font.size = Pt(12)
    
    # Add Table
    table = doc.add_table(rows=1, cols=4)
    table.style = 'Table Grid'
    
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'DATE'
    hdr_cells[1].text = 'TASK'
    hdr_cells[2].text = 'DELIVERABLE'
    hdr_cells[3].text = 'STATUS'
    
    # Make header bold
    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            for run in paragraph.runs:
                run.bold = True
    
    # August tasks based on actual work
    tasks = [
        ("August 3rd, 2026", "Product Upload", "Uploaded new products with names, descriptions, specifications, pricing and images", "Completed"),
        ("August 3rd, 2026", "Jumia Store Onboarding", "Continued preparing Jumia onboarding documentation and store setup requirements", "Ongoing"),
        
        ("August 5th, 2026", "Referral DB Schema Update", "Updated Users and Orders collections schema in Payload CMS to include payout fields (bank details) and referral tracking", "Completed"),
        ("August 5th, 2026", "Konga KYC & Registration", "Followed up on KYC verification status for the Konga store", "Ongoing"),
        
        ("August 7th, 2026", "Product Upload (Continued)", "Uploaded additional new arrivals to the website", "Completed"),
        ("August 7th, 2026", "CDCare Partnership", "Continued engagement and partnership discussions with CDCare", "Ongoing"),
        
        ("August 10th, 2026", "Checkout Integration", "Implemented referral code input field on the Checkout page to capture and validate user-applied codes", "Completed"),
        
        ("August 12th, 2026", "Commission Hooks", "Updated the Orders collection to save applied referral codes and trigger automatic commission calculation hooks", "Completed"),
        
        ("August 14th, 2026", "Admin UI Components", "Developed custom admin components (ReferralAmount, ExpectedPayoutField) for the Payload CMS admin panel", "Completed"),
        ("August 14th, 2026", "Product Description Update", "Reviewed and updated product descriptions, specifications and images across categories", "Completed"),
        
        ("August 17th, 2026", "Admin UI Debugging", "Fixed UI-level crashes in custom admin components to stabilize product editing and prevent form state errors", "Completed"),
        ("August 17th, 2026", "CDCare Partnership", "Drafting terms and continuing marketplace listing negotiations with CDCare", "Ongoing"),
        
        ("August 19th, 2026", "Schema Troubleshooting", "Resolved 'no such column: referrer_id' SQLite errors and reconciled authentication strategy conflicts", "Completed"),
        
        ("August 21st, 2026", "Product Upload", "Uploaded a new batch of products and verified their display on the storefront", "Completed"),
        ("August 21st, 2026", "Product Visibility Review", "Reviewed uploaded products to ensure consistent display and accurate pricing", "Completed"),
        
        ("August 24th, 2026", "Upload Performance Fix", "Diagnosed and resolved 'MaxListenersExceededWarning' memory leak issues during product image uploads", "Completed"),
        ("August 24th, 2026", "Konga KYC & Registration", "Resubmitted updated KYC verification documentation for Konga store", "Ongoing"),
        
        ("August 26th, 2026", "Referral System Verification", "Verified the entire end-to-end referral workflow to ensure earnings are correctly recorded and processed without database errors", "Completed"),
        ("August 26th, 2026", "Jumia Store Onboarding", "Followed up on marketplace approval for the Jumia store", "Ongoing"),
        
        ("August 28th, 2026", "Influencer Content Creation", "Created promotional content and guidelines for influencers to use with the referral engine", "Completed"),
        ("August 28th, 2026", "Influencer Content Distribution", "Shared the created promotional content with key influencers for the referral program", "Completed"),
        
        ("August 31st, 2026", "Influencer Content Updates", "Updated referral content and marketing assets based on early feedback from influencers", "Completed")
    ]
    
    for date, task, deliverable, status in tasks:
        row_cells = table.add_row().cells
        row_cells[0].text = date
        row_cells[1].text = task
        row_cells[2].text = deliverable
        row_cells[3].text = status
        
    doc.save('Reports/Nurudeen-Akindele-August.docx')
    print("Successfully created Reports/Nurudeen-Akindele-August.docx")

if __name__ == '__main__':
    create_august_report()
