import docx
import sys

def read_docx(file_path):
    doc = docx.Document(file_path)
    
    print("--- PARAGRAPHS ---")
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip():
            print(f"P{i}: {para.text}")
            
    print("\n--- TABLES ---")
    for t, table in enumerate(doc.tables):
        print(f"Table {t}:")
        for r, row in enumerate(table.rows):
            row_data = []
            for cell in row.cells:
                row_data.append(cell.text.replace("\n", " ").strip())
            print(f"  Row {r}: {' | '.join(row_data)}")

if __name__ == "__main__":
    read_docx(sys.argv[1])
